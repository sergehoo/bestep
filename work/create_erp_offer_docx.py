from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SKILL_DIR = Path('/Users/ogahserge/.codex/plugins/cache/openai-primary-runtime/documents/26.715.12143/skills/documents')
sys.path.insert(0, str(SKILL_DIR / 'scripts'))
from table_geometry import apply_table_geometry

OUT = Path('/Users/ogahserge/Documents/best_epargne/output/Offre_ERP_OLIVE_DANAROY.docx')
LOGO = Path('/Users/ogahserge/Documents/best_epargne/work/source_unpacked/ppt/media/image5.png')
COVER = Path('/Users/ogahserge/Documents/best_epargne/work/source_unpacked/ppt/media/image2.jpg')

# Preset: narrative_proposal. Named overrides: OLIVE DANAROY brand green/olive palette.
GREEN = '234E2B'
GREEN_2 = '397443'
OLIVE = '9DB43B'
YELLOW = 'E9D928'
INK = '17251C'
SLATE = '526159'
PALE = 'E6EFE7'
MIST = 'F3F7F3'
WHITE = 'FFFFFF'
LINE = 'D8E3D9'
AMBER = 'D69A22'

doc = Document()
doc.core_properties.title = 'Offre technique et financière — ERP OLIVE DANAROY'
doc.core_properties.subject = 'Conception et déploiement d’un ERP de gestion intégré'
doc.core_properties.author = 'Proposition préparée pour OLIVE DANAROY SARLU'
doc.core_properties.keywords = 'ERP, OLIVE DANAROY, BTP, restaurant, salles événementielles, lavage automobile'
doc.core_properties.comments = 'Document de travail — juillet 2026'


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name='Calibri', size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_borders(cell, color=LINE, size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        elem = tc_borders.find(tag)
        if elem is None:
            elem = OxmlElement(f'w:{edge}')
            tc_borders.append(elem)
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), size)
        elem.set(qn('w:color'), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        p_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def paragraph_left_border(paragraph, color=OLIVE, size='18', space='10'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), space)
    left.set(qn('w:color'), color)
    p_bdr.append(left)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    el = p_pr.find(qn('w:keepNext'))
    if el is None:
        el = OxmlElement('w:keepNext')
        p_pr.append(el)


def keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    el = p_pr.find(qn('w:keepLines'))
    if el is None:
        el = OxmlElement('w:keepLines')
        p_pr.append(el)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=SLATE)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn('w:pgNumType'))
    if pg_num is None:
        pg_num = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num)
    pg_num.set(qn('w:start'), str(start))


def configure_section(section, cover=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8 if cover else 1.0)
    section.bottom_margin = Inches(0.75 if cover else 1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def setup_styles():
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles['Title']
    title.font.name = 'Calibri'
    title._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    title._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(INK)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles['Subtitle']
    subtitle.font.name = 'Calibri'
    subtitle._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    subtitle._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(SLATE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    h1._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(GREEN)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    h2._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(GREEN_2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    h3._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(GREEN)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = styles['Caption']
    caption.font.name = 'Calibri'
    caption._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    caption._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(SLATE)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def create_numbering(num_fmt='bullet', level_text='•'):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'singleLevel')
    abstract.append(multi)
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)
    fmt = OxmlElement('w:numFmt')
    fmt.set(qn('w:val'), num_fmt)
    lvl.append(fmt)
    text = OxmlElement('w:lvlText')
    text.set(qn('w:val'), level_text)
    lvl.append(text)
    jc = OxmlElement('w:lvlJc')
    jc.set(qn('w:val'), 'left')
    lvl.append(jc)
    p_pr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'num')
    tab.set(qn('w:pos'), '540')
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '540')
    ind.set(qn('w:hanging'), '280')
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement('w:rPr')
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), 'Calibri')
    r_fonts.set(qn('w:hAnsi'), 'Calibri')
    r_pr.append(r_fonts)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), GREEN_2)
    r_pr.append(color)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abs_id = OxmlElement('w:abstractNumId')
    abs_id.set(qn('w:val'), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


setup_styles()
BULLET_NUM_ID = create_numbering('bullet', '•')
DECIMAL_NUM_ID = create_numbering('decimal', '%1.')


def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    keep_with_next(p)
    return p


def add_body(text, bold_lead=None, italic=False, align=None, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK, italic=italic)
    keep_together(p)
    return p


def add_bullet(text, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), str(BULLET_NUM_ID))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    if lead and text.startswith(lead):
        a = p.add_run(lead)
        set_run_font(a, size=11, color=INK, bold=True)
        b = p.add_run(text[len(lead):])
        set_run_font(b, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_numbered(text, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), str(DECIMAL_NUM_ID))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    if lead and text.startswith(lead):
        a = p.add_run(lead)
        set_run_font(a, size=11, color=INK, bold=True)
        b = p.add_run(text[len(lead):])
        set_run_font(b, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_callout(label, text, fill=PALE, accent=OLIVE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    paragraph_shading(p, fill)
    paragraph_left_border(p, accent, '18', '10')
    a = p.add_run(label.upper() + '  ')
    set_run_font(a, size=10, color=GREEN, bold=True)
    b = p.add_run(text)
    set_run_font(b, size=10.5, color=INK, bold=False)
    keep_together(p)
    return p


def add_table(rows, widths, header=True, right_cols=None, font_size=9.5, zebra=True):
    table = doc.add_table(rows=0, cols=len(widths))
    table.autofit = False
    right_cols = set(right_cols or [])
    for r_i, row_data in enumerate(rows):
        cells = table.add_row().cells
        for c_i, value in enumerate(row_data):
            cell = cells[c_i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if r_i == 0 and header:
                set_cell_shading(cell, GREEN)
            elif zebra and r_i % 2 == 1:
                set_cell_shading(cell, MIST)
            else:
                set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_i in right_cols else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            is_header_row = header and r_i == 0
            text_color = YELLOW if (is_header_row and c_i in right_cols) else WHITE if is_header_row else GREEN if c_i in right_cols else INK
            set_run_font(run, size=font_size, color=text_color, bold=(is_header_row or c_i in right_cols))
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120, cell_margins_dxa={'top': 80, 'bottom': 80, 'start': 120, 'end': 120})
    if header:
        set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_toc():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    begin.set(qn('w:dirty'), 'true')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = 'Le sommaire sera mis à jour automatiquement à l’ouverture dans Microsoft Word.'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, size=10.5, color=SLATE, italic=True)


def add_section_intro(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=12, color=SLATE, italic=True)
    keep_together(p)
    return p


def format_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run('OFFRE ERP  |  OLIVE DANAROY')
    set_run_font(hr, size=8.5, color=GREEN, bold=True)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run('JUILLET 2026   •   ')
    set_run_font(label, size=8.5, color=SLATE)
    add_page_number(fp)


# Cover section
cover_section = doc.sections[0]
configure_section(cover_section, cover=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
p.add_run().add_picture(str(LOGO), width=Inches(1.8))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('PROPOSITION TECHNIQUE & FINANCIÈRE')
set_run_font(r, size=10, color=GREEN_2, bold=True)

p = doc.add_paragraph(style='Title')
p.add_run('ERP DE GESTION INTÉGRÉ')

p = doc.add_paragraph(style='Subtitle')
p.add_run('Une plateforme unique pour piloter les projets, les ventes, les opérations et la rentabilité de toutes les activités d’OLIVE DANAROY.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run('PRÉPARÉ POUR  OLIVE DANAROY SARLU\n16 juillet 2026 — Abidjan, Côte d’Ivoire')
set_run_font(r, size=10.5, color=SLATE, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
p.add_run().add_picture(str(COVER), width=Inches(6.3))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DOCUMENT DE TRAVAIL  •  VALIDITÉ 30 JOURS')
set_run_font(r, size=8.5, color=GREEN, bold=True)

# Main section
main = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(main, cover=False)
set_page_number_start(main, 1)
format_header_footer(main)

doc.add_paragraph('Sommaire', style='Heading 1')
add_body('Les titres sont structurés avec les styles Word afin de faciliter la navigation dans le document.', italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
toc_rows = [
    ['1. Synthèse exécutive', '2'],
    ['2. Contexte, enjeux et périmètre', '3'],
    ['3. Offre fonctionnelle', '4'],
    ['4. Architecture technique, sécurité et continuité', '7'],
    ['5. Tableaux de bord et indicateurs', '8'],
    ['6. Méthodologie et planning de réalisation', '9'],
    ['7. Gouvernance, formation et support', '10'],
    ['8. Offre financière', '11'],
    ['9. Hypothèses, exclusions et conditions', '13'],
    ['10. Prochaines étapes et bon pour accord', '14']
]
toc_table = add_table(toc_rows, [8400, 960], header=False, right_cols=[1], font_size=10.2, zebra=True)
for row in toc_table.rows:
    for cell in row.cells:
        set_cell_borders(cell, color='E4ECE5', size='2')
add_callout('Repères clés', 'Budget projet : 36 850 000 FCFA HT  •  Délai : 24 semaines  •  Dimensionnement initial : 50 utilisateurs et 5 sites  •  Hypercare : 4 semaines.', fill='EFF4DA', accent=GREEN_2)

# 1
add_heading('1. Synthèse exécutive', 1)
add_section_intro('OLIVE DANAROY a besoin d’un système unique capable d’accompagner la diversité de ses métiers sans multiplier les outils ni les doubles saisies.')
add_callout('Recommandation', 'Déployer un ERP modulaire autour d’un socle commun — ventes, achats, stocks, finance, RH, documents et tableaux de bord — puis activer les fonctions propres au BTP, au transport, au restaurant, aux salles événementielles, au lavage automobile, au traiteur et aux résidences.')
add_body('L’objectif est de remplacer les fichiers dispersés et les suivis manuels par une base de données centralisée, des circuits de validation, une traçabilité complète et une lecture fiable de la rentabilité par activité, site, projet et client.')

metric_rows = [
    ['24 SEMAINES', '50 UTILISATEURS', '5 SITES', '4 SEMAINES'],
    ['Cadrage à mise en production', 'Hypothèse initiale', 'Dimensionnement initial', 'Hypercare inclus']
]
metrics = add_table(metric_rows, [2340, 2340, 2340, 2340], header=True, font_size=9.2, zebra=False)
for cell in metrics.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = rgb(YELLOW)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for cell in metrics.rows[1].cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_callout('Budget indicatif', '36 850 000 FCFA HT pour la conception, le développement, la migration, la formation, le déploiement et quatre semaines d’accompagnement renforcé. Les frais récurrents et services tiers sont présentés séparément.', fill='EFF4DA', accent=GREEN_2)

# 2
add_heading('2. Contexte, enjeux et périmètre', 1)
add_section_intro('Le dossier technique de l’entreprise présente une organisation multisectorielle : BTP, communication, informatique, transport, bureautique et services divers.')
add_body('La solution proposée couvre les domaines existants et les nouvelles activités explicitement demandées : gestion du restaurant, location de salles événementielles et lavage automobile. Le traiteur et la location de résidences, également mentionnés dans le dossier, sont intégrés au périmètre fonctionnel.')
add_heading('2.1 Enjeux de gestion', 2)
for item, lead in [
    ('Centraliser les référentiels clients, fournisseurs, articles, collaborateurs, caisses, banques et centres de coûts.', 'Centraliser'),
    ('Fiabiliser les validations, les pièces justificatives, les encaissements, les dépenses et les stocks.', 'Fiabiliser'),
    ('Mesurer la rentabilité par chantier, commande, trajet, événement, salle, plat, véhicule et site.', 'Mesurer'),
    ('Accélérer la facturation, le recouvrement, les clôtures de caisse et la production de rapports.', 'Accélérer'),
    ('Sécuriser les accès, les sauvegardes, les responsabilités et la piste d’audit.', 'Sécuriser')
]:
    add_bullet(item, lead)

add_heading('2.2 Activités couvertes', 2)
activities = [
    'Bâtiments et travaux publics : construction, rénovation, réhabilitation, démolition, voirie et canalisation.',
    'Communication : conception, branding, enseignes, signalétique et impression.',
    'Informatique : matériels, consommables, logiciels métiers, réseaux, VPN et vidéosurveillance.',
    'Bureautique : équipements, mobilier, fournitures et aménagement de bureaux.',
    'Transport et logistique : enlèvement, livraison, location de véhicules et gestion de flotte.',
    'Restaurant et cafétéria : vente, cuisine, recettes, stock, réservations et caisse.',
    'Salles événementielles : disponibilités, devis, contrats, acomptes et prestations.',
    'Lavage automobile : tickets, forfaits, abonnements, postes, agents et caisse.',
    'Traiteur : devis, menus, besoins matières, production et marge par événement.',
    'Résidences : disponibilités, séjours, dépôts de garantie et facturation.'
]
for a in activities:
    add_bullet(a)

# 3
add_heading('3. Offre fonctionnelle', 1)
add_section_intro('Le principe directeur est simple : une donnée saisie une fois doit alimenter automatiquement les opérations, la trésorerie, la comptabilité et les tableaux de bord.')

add_heading('3.1 Socle transversal de l’ERP', 2)
core = [
    ('CRM et ventes : ', 'prospects, opportunités, devis, commandes, contrats, factures, règlements et relances.'),
    ('Achats : ', 'demandes d’achat, validations, commandes fournisseurs, réceptions et contrôle des factures.'),
    ('Stocks : ', 'multi-dépôts, lots, numéros de série, mouvements, inventaires, seuils et valorisation.'),
    ('Finance et trésorerie : ', 'caisses, banques, créances, dettes, dépenses, centres de coûts et prévisions.'),
    ('Ressources humaines : ', 'personnel, présence, congés, feuilles de temps et préparation des éléments de paie.'),
    ('Documents et validations : ', 'contrats, pièces jointes, numérotation, visas, alertes et journal d’audit.'),
    ('Portail et mobile : ', 'accès web responsive, PWA mobile, notifications et validations à distance.'),
    ('Pilotage : ', 'budgets, marges, objectifs, tableaux de bord et exports Excel/PDF.')
]
for lead, detail in core:
    add_bullet(lead + detail, lead)

add_heading('3.2 Module BTP et gestion de chantiers', 2)
add_body('Le module BTP suit le cycle complet, de l’appel d’offres à la marge finale : estimation, planification, exécution, situations de travaux, réception et clôture.')
for lead, detail in [
    ('Études et devis : ', 'estimations, bordereaux, versions, validation et transformation en projet.'),
    ('Planification : ', 'WBS, tâches, jalons, ressources, équipes, engins et matériels.'),
    ('Contrôle des coûts : ', 'budgets par lot, achats, sous-traitance, consommations, temps et écarts.'),
    ('Suivi terrain : ', 'avancement, photos, incidents, qualité, sécurité, réserves et rapports.'),
    ('Facturation : ', 'attachements, situations, acomptes, retenues, règlements et recouvrement.'),
    ('Rentabilité : ', 'prévu, engagé, réalisé, reste à faire et marge à terminaison.')
]:
    add_bullet(lead + detail, lead)

add_heading('3.3 Commerce, informatique, communication et bureautique', 2)
add_body('Ces activités utilisent un cycle commun devis–commande–approvisionnement–livraison–facturation–encaissement, enrichi par des fonctions spécifiques aux équipements et prestations.')
for item in [
    'Catalogues, tarifs, remises, kits, numéros de série, garanties et livraisons partielles.',
    'Contrats de maintenance, abonnements, tickets d’assistance, temps passé, SLA et renouvellements.',
    'Briefs, bons à tirer, plan de charge créatif, impression, sous-traitance et marge par campagne.',
    'Stock de matériels, mobilier et consommables avec inventaires et seuils d’alerte.'
]:
    add_bullet(item)

add_heading('3.4 Transport, logistique, résidences et traiteur', 2)
add_body('Le transport et les services associés sont planifiés et facturés depuis la même base clients, fournisseurs et trésorerie.')
for lead, detail in [
    ('Transport et flotte : ', 'ordres de mission, véhicules, chauffeurs, documents, carburant, péages, avances, entretien, preuve de livraison et coût par kilomètre.'),
    ('Résidences : ', 'disponibilités, réservation, check-in/check-out, dépôt de garantie, services annexes, état des chambres et facturation.'),
    ('Traiteur : ', 'devis par nombre de couverts, menus, fiches techniques, besoins matières, planning de production, équipes, livraison et marge par événement.')
]:
    add_bullet(lead + detail, lead)

add_heading('3.5 Module restaurant', 2)
add_body('Le module restaurant relie le point de vente à la cuisine, aux recettes, au stock et à la comptabilité. Il doit rester rapide en caisse tout en garantissant une traçabilité complète.')
for lead, detail in [
    ('Point de vente : ', 'tables, comptoir, emporter, livraison, remises, annulations, paiements multiples et clôture de caisse.'),
    ('Cuisine : ', 'tickets, priorités, statuts, temps de préparation, ruptures et contrôle de sortie.'),
    ('Recettes et stock : ', 'fiches techniques, portions, consommations automatiques, pertes, transferts et inventaires.'),
    ('Clients : ', 'réservations, historique, fidélité, comptes entreprises et abonnements.'),
    ('Indicateurs : ', 'chiffre d’affaires journalier, ticket moyen, coût matière, pertes, marge par plat et écart de caisse.')
]:
    add_bullet(lead + detail, lead)

add_heading('3.6 Module location de salles événementielles', 2)
add_body('Le module sécurise les disponibilités, évite les doubles réservations et rend visible le chiffre d’affaires prévisionnel.')
for lead, detail in [
    ('Calendrier : ', 'salles, capacités, créneaux, options, confirmations et dates d’expiration.'),
    ('Devis et contrat : ', 'salle, décoration, sonorisation, restauration, personnel, équipements et conditions.'),
    ('Acomptes et solde : ', 'échéanciers, reçus, relances, dépôt de garantie et remboursement.'),
    ('Exécution : ', 'checklists, plan de salle, affectations, incidents et état des lieux.'),
    ('Pilotage : ', 'taux d’occupation, recettes futures, acomptes, impayés et marge par événement.')
]:
    add_bullet(lead + detail, lead)

add_heading('3.7 Module lavage automobile', 2)
add_body('Chaque véhicule reçoit un ticket qui suit les étapes arrivée, diagnostic, choix du forfait, traitement, contrôle qualité et paiement.')
for lead, detail in [
    ('Offres : ', 'lavage simple, complet, VIP, options et tarifs selon la catégorie du véhicule.'),
    ('Abonnements : ', 'cartes prépayées, comptes entreprises, flottes et consommation par véhicule.'),
    ('Opérations : ', 'file d’attente, affectation des postes, agents, temps de cycle et reprises.'),
    ('Caisse et contrôle : ', 'paiements, remises, annulations, consommables, incidents et clôture.'),
    ('Indicateurs : ', 'véhicules par jour, temps moyen, revenu par poste, consommables par véhicule et écart de caisse.')
]:
    add_bullet(lead + detail, lead)

# 4
add_heading('4. Architecture technique, sécurité et continuité', 1)
add_section_intro('La solution cible est une application web responsive, centralisée, sécurisée et conçue pour évoluer par modules.')
add_heading('4.1 Architecture cible', 2)
for lead, detail in [
    ('Couche utilisateurs : ', 'direction, finance, ventes, chantiers, transport, restaurant, caisses et accès mobile.'),
    ('Couche applicative : ', 'modules métiers, workflows, notifications, tableaux de bord et portail.'),
    ('Services : ', 'API, génération de documents, e-mail/SMS, paiements, imports/exports et journalisation.'),
    ('Données et cloud : ', 'base centralisée, sauvegardes, supervision, chiffrement et mécanismes de reprise.')
]:
    add_bullet(lead + detail, lead)

add_heading('4.2 Sécurité', 2)
for item in [
    'Gestion des profils, rôles, droits par activité et seuils d’autorisation.',
    'Chiffrement des échanges HTTPS, protection des secrets et séparation des environnements.',
    'Sauvegardes quotidiennes avec rétention, restauration testée et copie externalisée.',
    'Journal des connexions, validations, modifications sensibles, annulations et exports.',
    'Double authentification disponible en option selon les profils et risques.'
]:
    add_bullet(item)
add_callout('Point à valider', 'Les objectifs de disponibilité, les durées de conservation, le plan de reprise et les règles de protection des données seront précisés lors du cadrage et repris dans le dossier d’architecture.')

# 5
add_heading('5. Tableaux de bord et indicateurs', 1)
add_section_intro('La direction dispose d’une vue consolidée et peut descendre jusqu’à la transaction source pour expliquer chaque chiffre.')
for lead, detail in [
    ('Direction : ', 'chiffre d’affaires, marge, trésorerie, créances, dettes, budgets et tendances par activité.'),
    ('BTP : ', 'avancement, engagé, réalisé, écarts, retards, facturation et marge à terminaison.'),
    ('Restaurant : ', 'ticket moyen, coût matière, pertes, marge par plat, productivité et écart de caisse.'),
    ('Salles : ', 'taux d’occupation, chiffre d’affaires futur, acomptes, soldes et marge par événement.'),
    ('Lavage : ', 'véhicules par jour, temps de cycle, revenu par poste, reprises et consommables.'),
    ('Transport : ', 'coût au kilomètre, carburant, disponibilité, immobilisation et marge par trajet.'),
    ('Commerce : ', 'pipeline, conversion, marge par article, rotation des stocks et renouvellements.')
]:
    add_bullet(lead + detail, lead)

# 6
add_heading('6. Méthodologie et planning de réalisation', 1)
add_section_intro('Le projet est réalisé de manière incrémentale, avec une démonstration toutes les deux semaines et une validation formelle à chaque jalon.')
schedule_rows = [
    ['Phase', 'Période', 'Objet', 'Livrable principal'],
    ['1', 'Semaines 1 à 3', 'Cadrage et blueprint', 'Processus cibles, règles, périmètre et backlog validés'],
    ['2', 'Semaines 4 à 6', 'UX/UI, prototype et architecture', 'Parcours clés et architecture approuvés'],
    ['3', 'Semaines 7 à 12', 'Socle ERP et module BTP', 'Version intégrée disponible en recette'],
    ['4', 'Semaines 13 à 17', 'Modules métiers', 'Restaurant, salles, lavage, transport et services en recette'],
    ['5', 'Semaines 18 à 20', 'Migration et intégrations', 'Données reprises et interfaces vérifiées'],
    ['6', 'Semaines 21 à 22', 'Recette et formation', 'Procès-verbal de recette fonctionnelle'],
    ['7', 'Semaines 23 à 24', 'Mise en production', 'Démarrage, sauvegardes et support actif'],
    ['8', '4 semaines après démarrage', 'Hypercare', 'Stabilisation et accompagnement renforcé']
]
add_table(schedule_rows, [900, 1900, 2600, 3960], header=True, font_size=8.7, zebra=True)

add_heading('6.1 Principaux livrables', 2)
for item in [
    'Dossier de conception fonctionnelle et cartographie des processus cibles.',
    'Prototype validé et dossier d’architecture technique.',
    'Applications configurées, code source, scripts de déploiement et documentation.',
    'Plan de migration, données reprises et rapport de contrôle.',
    'Cahier de recette, résultats des tests et procès-verbal de validation.',
    'Guides utilisateurs, procédures d’exploitation et supports de formation.',
    'Plan de sauvegarde, supervision et procédure de support.'
]:
    add_bullet(item)

# 7
add_heading('7. Gouvernance, formation et support', 1)
add_heading('7.1 Gouvernance', 2)
for lead, detail in [
    ('Comité de pilotage mensuel : ', 'direction générale, sponsor, chef de projet client et chef de projet intégrateur.'),
    ('Point projet hebdomadaire : ', 'avancement, décisions, actions, risques, données et tests.'),
    ('Référents métiers : ', 'BTP, finance, ventes, transport, restaurant, salles, lavage, traiteur et résidences.'),
    ('Démonstration bimensuelle : ', 'validation progressive des fonctions et collecte des retours.')
]:
    add_bullet(lead + detail, lead)

add_heading('7.2 Formation et transfert', 2)
training_rows = [
    ['Public', 'Durée indicative', 'Contenu'],
    ['Administrateurs', '2 jours', 'Paramétrage, profils, référentiels, sauvegardes et support de niveau 1'],
    ['Super-utilisateurs', '3 jours', 'Processus métiers, contrôles, recette et accompagnement des équipes'],
    ['Utilisateurs finaux', '6 sessions', 'Formation par rôle et exercices sur données de démonstration'],
    ['Direction', '½ journée', 'Tableaux de bord, alertes, analyse et exports']
]
add_table(training_rows, [2300, 1800, 5260], header=True, font_size=9.0, zebra=True)

add_heading('7.3 Support', 2)
add_body('Quatre semaines d’hypercare sont incluses après la mise en production. Cette période couvre la stabilisation, l’accompagnement des utilisateurs, la correction des anomalies et le suivi rapproché des opérations critiques.')

# 8
add_heading('8. Offre financière', 1)
add_section_intro('Les montants sont exprimés en francs CFA hors taxes. Le prix définitif sera confirmé après l’atelier de cadrage et la validation du périmètre contractuel.')
budget_rows = [
    ['Poste', 'Montant HT'],
    ['Cadrage, cartographie des processus et blueprint', '2 250 000'],
    ['UX/UI, prototype et architecture technique', '2 500 000'],
    ['Socle ERP : CRM, ventes, achats, stocks, finance, RH et documents', '9 800 000'],
    ['Module BTP et gestion de projets / chantiers', '4 600 000'],
    ['Module transport, logistique et flotte', '3 200 000'],
    ['Module restaurant : POS, cuisine, recettes et stock', '3 300 000'],
    ['Modules salles, traiteur et résidences', '2 600 000'],
    ['Module lavage automobile', '1 900 000'],
    ['Reporting, portail et expérience mobile PWA', '2 100 000'],
    ['Migration, paramétrage et intégrations standards', '2 400 000'],
    ['Recette, formation, déploiement et hypercare', '2 200 000'],
    ['TOTAL PROJET HT', '36 850 000 FCFA']
]
budget_table = add_table(budget_rows, [7440, 1920], header=True, right_cols=[1], font_size=9.1, zebra=True)
for cell in budget_table.rows[-1].cells:
    set_cell_shading(cell, INK)
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = rgb(YELLOW if cell == budget_table.rows[-1].cells[0] else WHITE)

add_body('TVA et retenues éventuelles selon la réglementation applicable. Les services tiers et les frais de déplacement hors Abidjan ne sont pas inclus.', italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=10)

add_heading('8.1 Échéancier de paiement proposé', 2)
payment_rows = [
    ['Échéance', 'Jalon', 'Montant HT'],
    ['20 %', 'Commande et lancement', '7 370 000 FCFA'],
    ['20 %', 'Blueprint et prototype validés', '7 370 000 FCFA'],
    ['25 %', 'Socle ERP et BTP disponibles en recette', '9 212 500 FCFA'],
    ['20 %', 'Modules métiers disponibles en recette', '7 370 000 FCFA'],
    ['10 %', 'Mise en production', '3 685 000 FCFA'],
    ['5 %', 'Réception après hypercare', '1 842 500 FCFA']
]
add_table(payment_rows, [1200, 5460, 2700], header=True, right_cols=[2], font_size=9.1, zebra=True)

add_heading('8.2 Coûts récurrents indicatifs — non inclus', 2)
recurring_rows = [
    ['Service', 'Montant mensuel HT'],
    ['Cloud, sauvegardes, certificats et supervision', '240 000 FCFA'],
    ['Support correctif et mises à jour', '350 000 FCFA'],
    ['TOTAL RÉCURRENT INDICATIF', '590 000 FCFA / mois']
]
rec_table = add_table(recurring_rows, [6600, 2760], header=True, right_cols=[1], font_size=9.2, zebra=True)
for cell in rec_table.rows[-1].cells:
    set_cell_shading(cell, PALE)
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = rgb(GREEN)
add_body('Les SMS, WhatsApp, paiements en ligne, équipements de caisse, imprimantes, lecteurs et prestations évolutives seront facturés selon consommation ou devis.', italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

# 9
add_heading('9. Hypothèses, exclusions et conditions', 1)
add_heading('9.1 Hypothèses structurantes', 2)
for item in [
    'Jusqu’à 50 utilisateurs nommés et 5 sites opérationnels dans le périmètre initial.',
    'Données sources fournies en fichiers Excel/CSV structurés, nettoyés et validés par le client.',
    'Un référent disponible par activité pour les ateliers, les validations et les tests.',
    'Interfaces tierces limitées aux API documentées et accessibles pendant le projet.',
    'Plan comptable, règles de gestion, modèles de documents et circuits de validation fournis par le client.',
    'Environnements de test et de production hébergés selon l’offre cloud retenue.'
]:
    add_bullet(item)

add_heading('9.2 Éléments à confirmer au cadrage', 2)
for item in [
    'Nombre exact de sites, caisses, dépôts, salles, postes de lavage et terminaux.',
    'Périmètre comptable, fiscal et paie ainsi que les exigences réglementaires associées.',
    'Historique à migrer, volumétrie, qualité des données et règles de rapprochement.',
    'Matériel POS, imprimantes cuisine, lecteurs, terminaux mobiles et réseau local.',
    'Intégrations bancaires, mobile money, paiement en ligne, SMS, WhatsApp et services externes.',
    'Objectifs de disponibilité, reprise après incident et durée de conservation des données.'
]:
    add_bullet(item)

add_heading('9.3 Exclusions indicatives', 2)
for item in [
    'Acquisition de matériels, câblage, terminaux, imprimantes et équipements réseaux.',
    'Licences ou abonnements de services tiers non expressément inclus.',
    'Nettoyage lourd ou reconstitution des données historiques non disponibles.',
    'Développements supplémentaires demandés après validation du périmètre, traités par avenant.',
    'Déplacements, hébergement et indemnités hors Abidjan, facturés après accord.'
]:
    add_bullet(item)

add_heading('9.4 Conditions commerciales', 2)
for item in [
    'Validité de l’offre : 30 jours à compter du 16 juillet 2026.',
    'Démarrage prévisionnel : après signature, paiement de l’acompte et disponibilité des référents.',
    'Les délais sont conditionnés par la validation rapide des livrables et la disponibilité des données.',
    'Toute modification substantielle du périmètre, des interfaces ou de la volumétrie fera l’objet d’un avenant.'
]:
    add_bullet(item)

# 10
add_heading('10. Prochaines étapes et bon pour accord', 1)
add_section_intro('La validation du principe permet d’organiser l’atelier de cadrage, de confirmer les hypothèses et d’émettre l’offre contractuelle finale.')
for step, lead in [
    ('Valider le principe et nommer le sponsor ainsi que le chef de projet client.', 'Valider'),
    ('Organiser un atelier de cadrage avec les responsables de chaque activité.', 'Organiser'),
    ('Confirmer les utilisateurs, sites, volumes, matériels et intégrations.', 'Confirmer'),
    ('Finaliser le périmètre, le planning, le prix et les clauses contractuelles.', 'Finaliser'),
    ('Signer l’accord et lancer la phase de cadrage.', 'Signer')
]:
    add_numbered(step, lead)

add_callout('Contact OLIVE DANAROY', 'E-mail : olivedanaroy@gmail.com  •  Téléphone : (+225) 07 77 99 99 48  •  Siège : Abidjan, Cocody, Angré Château.', fill='EFF4DA', accent=GREEN_2)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
sig_rows = [
    ['POUR OLIVE DANAROY SARLU', 'POUR LE PRESTATAIRE / INTÉGRATEUR'],
    ['Nom et fonction :\n\nDate :\n\nSignature et cachet :\n\n\n', 'Nom et fonction :\n\nDate :\n\nSignature et cachet :\n\n\n']
]
sig_table = add_table(sig_rows, [4680, 4680], header=True, font_size=9.3, zebra=False)
for cell in sig_table.rows[1].cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

add_body('La signature de cette page vaut accord de principe sur les orientations de l’offre. Le contrat final précisera les engagements, responsabilités, niveaux de service et conditions juridiques.', italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=0)

# Ask Word/LibreOffice to refresh fields.
settings = doc.settings._element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
